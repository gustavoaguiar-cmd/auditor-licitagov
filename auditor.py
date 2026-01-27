import streamlit as st
import os
import psycopg2
from datetime import datetime
from io import BytesIO
from docx import Document as DocxDocument
from pypdf import PdfReader
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.prompts import PromptTemplate
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from dotenv import load_dotenv

# Carrega variáveis de ambiente
load_dotenv()

# --- CONFIGURAÇÃO DO BANCO DE DADOS (PostgreSQL) ---
def get_db_connection():
    try:
        url = os.environ.get("DATABASE_URL")
        conn = psycopg2.connect(url)
        return conn
    except Exception as e:
        st.error(f"Erro ao conectar no Banco de Dados: {e}")
        return None

def init_db():
    conn = get_db_connection()
    if conn:
        cur = conn.cursor()
        # Tabelas essenciais
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                username VARCHAR(50) PRIMARY KEY,
                password VARCHAR(50) NOT NULL,
                role VARCHAR(20) DEFAULT 'user',
                perm_auditor BOOLEAN DEFAULT FALSE,
                perm_gerador BOOLEAN DEFAULT FALSE,
                perm_parecer BOOLEAN DEFAULT FALSE,
                perm_pca BOOLEAN DEFAULT FALSE,
                perm_recursos BOOLEAN DEFAULT FALSE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS system_logs (
                id SERIAL PRIMARY KEY,
                username VARCHAR(50),
                action VARCHAR(200),
                details TEXT,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS support_tickets (
                id SERIAL PRIMARY KEY,
                username VARCHAR(50),
                message TEXT,
                status VARCHAR(20) DEFAULT 'aberto',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
        """)
        # Cria ADMIN padrão se não existir
        cur.execute("SELECT * FROM users WHERE username = 'admin'")
        if not cur.fetchone():
            cur.execute("""
                INSERT INTO users (username, password, role, perm_auditor, perm_gerador, perm_parecer, perm_pca, perm_recursos)
                VALUES ('admin', 'admin123', 'admin', TRUE, TRUE, TRUE, TRUE, TRUE)
            """)
        conn.commit()
        cur.close()
        conn.close()

def log_action(username, action, details=""):
    conn = get_db_connection()
    if conn:
        cur = conn.cursor()
        cur.execute("INSERT INTO system_logs (username, action, details) VALUES (%s, %s, %s)", (username, action, details))
        conn.commit()
        conn.close()

def send_support_ticket(username, message):
    conn = get_db_connection()
    if conn:
        cur = conn.cursor()
        cur.execute("INSERT INTO support_tickets (username, message) VALUES (%s, %s)", (username, message))
        conn.commit()
        conn.close()
        st.success("✅ Chamado aberto com sucesso!")

# --- INICIALIZAÇÃO DB ---
if "db_initialized" not in st.session_state:
    init_db()
    st.session_state["db_initialized"] = True

# --- MOTOR DE INTELIGÊNCIA (CÉREBRO V15) ---
@st.cache_resource
def load_knowledge_base():
    """Carrega a base de conhecimento com cache em disco."""
    index_path = "faiss_index"
    folder_path = "data/legislacao" # Certifique-se que esta pasta existe no seu GitHub com os PDFs
    embeddings = OpenAIEmbeddings()

    # 1. Tenta carregar índice salvo
    if os.path.exists(index_path):
        try:
            return FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        except:
            pass

    # 2. Se não existir, cria do zero lendo TUDO
    docs = []
    if not os.path.exists(folder_path):
        return None

    # Varredura completa recursiva
    for root, dirs, files in os.walk(folder_path):
        for filename in files:
            if filename.lower().endswith(".pdf"):
                file_path = os.path.join(root, filename)
                try:
                    reader = PdfReader(file_path)
                    text = ""
                    for page in reader.pages:
                        if page.extract_text():
                            text += page.extract_text()
                    if text:
                        docs.append(Document(page_content=text, metadata={"source": filename}))
                except:
                    pass
    
    if not docs:
        return None

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    splits = text_splitter.split_documents(docs)
    vectorstore = FAISS.from_documents(splits, embeddings)
    vectorstore.save_local(index_path)
    return vectorstore

def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        try:
            pdf_reader = PdfReader(pdf)
            for page in pdf_reader.pages:
                if page.extract_text():
                    text += page.extract_text()
        except:
            pass
    return text

def create_word_docx(markdown_text):
    doc = DocxDocument()
    doc.add_heading('Lici Govtech - Relatório de Auditoria', 0)
    for line in markdown_text.split('\n'):
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line.replace('- ', '').replace('* ', ''), style='List Bullet')
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def get_autonomous_prompt(doc_type):
    # O PROMPT ORIGINAL V15
    return """
    Você é um Auditor Federal de Controle Externo Especialista (Nível TCU).
    
    SUA MISSÃO:
    Auditar o documento ({doc_type}) com base na Lei 14.133/21 e na JURISPRUDÊNCIA fornecida.

    REGRAS DE OURO:
    1. **LEGISLAÇÃO:** Priorize totalmente a Lei 14.133/2021.
    2. **CITAÇÕES:** PROCURE NO TEXTO DO CONTEXTO o número do Acórdão, Súmula ou Enunciado. Se não encontrar, cite genericamente "Jurisprudência TCU". NÃO cite nomes de arquivos PDF.
    3. **RIGOR:** Aponte riscos de sobrepreço, restrição de competitividade e direcionamento.

    ---
    CONTEXTO JURÍDICO (Base de Conhecimento):
    {context}
    ---

    DOCUMENTO A SER AUDITADO ({doc_type}):
    {text}

    GERE O RELATÓRIO NESTE FORMATO:

    ## 🚨 Relatório de Auditoria Técnica

    ### 1. Análise de Conformidade (Lei 14.133/21)
    (Análise geral do documento).

    ### 2. Riscos e Irregularidades Identificadas
    - **Ponto Crítico:** [Descreva o problema]
    - **Fundamentação:** [Cite o Acórdão X ou Artigo Y da Lei 14.133 do contexto]
    - **Recomendação:** [O que fazer]

    ### 3. Conclusão do Auditor
    """

# --- FRONTEND E NAVEGAÇÃO ---
st.set_page_config(page_title="Lici Govtech", page_icon="🏛️", layout="wide")

# CSS para esconder elementos padrão
st.markdown("""
<style>
    .stApp {background-color: #f8f9fa;}
    header {visibility: hidden;}
    footer {visibility: hidden;}
    .big-card {
        background-color: white; padding: 25px; border-radius: 12px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05); text-align: center;
        transition: transform 0.2s; cursor: pointer; border: 1px solid #e0e0e0;
        height: 200px; display: flex; flex-direction: column; justify-content: center;
    }
    .big-card:hover { transform: translateY(-5px); border-color: #0f2c4a; }
    h1, h2, h3 {color: #0f2c4a;}
</style>
""", unsafe_allow_html=True)

# Lógica de Login
if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

if not st.session_state["logged_in"]:
    col1, col2, col3 = st.columns([1,1,1])
    with col2:
        st.markdown("<br><br><h1 style='text-align: center;'>Lici Govtech 🏛️</h1>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: #666;'>Plataforma de Inteligência em Contratações</p>", unsafe_allow_html=True)
        
        usuario = st.text_input("Usuário")
        senha = st.text_input("Senha", type="password")
        
        if st.button("Acessar Plataforma", use_container_width=True):
            conn = get_db_connection()
            if conn:
                cur = conn.cursor()
                cur.execute("SELECT * FROM users WHERE username = %s AND password = %s", (usuario, senha))
                user_data = cur.fetchone()
                conn.close()
                
                if user_data:
                    st.session_state["logged_in"] = True
                    st.session_state["user"] = {
                        "username": user_data[0],
                        "role": user_data[2],
                        "perms": {
                            "Auditor": user_data[3],
                            "Gerador": user_data[4],
                            "Parecer": user_data[5],
                            "PCA": user_data[6],
                            "Recursos": user_data[7]
                        }
                    }
                    log_action(usuario, "LOGIN", "Login realizado")
                    st.rerun()
                else:
                    st.error("Credenciais inválidas.")
    st.stop()

# --- ÁREA LOGADA ---
user = st.session_state["user"]
api_key = os.environ.get("OPENAI_API_KEY")

# Sidebar
st.sidebar.title(f"Olá, {user['username']}")
menu = st.sidebar.radio("Menu", ["Dashboard", "Suporte / Ajuda"] + (["Admin"] if user["role"] == 'admin' else []))

if st.sidebar.button("Sair"):
    st.session_state["logged_in"] = False
    st.rerun()

# --- CARREGA CÉREBRO ---
with st.sidebar:
    st.markdown("---")
    st.caption("Status do Sistema:")
    with st.spinner("Conectando IA..."):
        vectorstore = load_knowledge_base()
    if vectorstore:
        st.success("✅ Base Jurídica Ativa")
    else:
        st.warning("⚠️ Base em construção")

# --- DASHBOARD ---
if menu == "Dashboard":
    
    # Estado do módulo
    if "modulo_ativo" not in st.session_state:
        st.session_state["modulo_ativo"] = None

    # Se nenhum módulo selecionado, mostra Menu
    if st.session_state["modulo_ativo"] is None:
        st.title("Painel Principal")
        st.write("Selecione uma ferramenta para iniciar:")
        
        # Grid Inteligente
        cols = st.columns(3)
        modulos = [
            ("🔍 Auditoria IA", "auditor", user["perms"]["Auditor"]),
            ("📄 Gerador Docs", "gerador", user["perms"]["Gerador"]),
            ("⚖️ Parecerista", "parecer", user["perms"]["Parecer"]),
            ("📅 PCA", "pca", user["perms"]["PCA"]),
            ("🛡️ Recursos", "recursos", user["perms"]["Recursos"])
        ]
        
        idx = 0
        for nome, codigo, permitido in modulos:
            if permitido:
                with cols[idx % 3]:
                    st.markdown(f'<div class="big-card"><h3>{nome}</h3></div>', unsafe_allow_html=True)
                    if st.button(f"Abrir {nome}", key=codigo, use_container_width=True):
                        st.session_state["modulo_ativo"] = codigo
                        st.rerun()
                idx += 1

    # --- MÓDULO AUDITOR (CÓDIGO V15 INTEGRADO AQUI) ---
    elif st.session_state["modulo_ativo"] == "auditor":
        st.button("⬅️ Voltar ao Painel", on_click=lambda: st.session_state.update({"modulo_ativo": None}))
        st.title("Auditoria Especializada 🔍")
        st.info("A IA analisará o documento cruzando com a Lei 14.133/21 e Jurisprudência.")

        doc_type = st.selectbox("Documento:", ["Edital de Licitação", "TR", "ETP", "Projeto Básico"])
        uploaded_file = st.file_uploader("Upload do Arquivo PDF", type="pdf")

        if uploaded_file and st.button("🚀 Iniciar Auditoria"):
            if not api_key:
                st.error("API Key não configurada.")
            else:
                with st.spinner("Lendo documento e consultando Base Jurídica..."):
                    try:
                        raw_text = get_pdf_text([uploaded_file])
                        if len(raw_text) < 50:
                            st.warning("⚠️ O PDF parece ser uma imagem digitalizada. O OCR será ativado na próxima versão.")
                        else:
                            # 1. Busca Contexto (RAG)
                            contexto = ""
                            if vectorstore:
                                docs_rel = vectorstore.similarity_search(raw_text[:6000], k=6)
                                for doc in docs_rel:
                                    contexto += f"\n[JURISPRUDÊNCIA]: {doc.page_content}\n"
                            
                            # 2. Chama LLM
                            llm = ChatOpenAI(model_name="gpt-4-turbo", temperature=0.2, openai_api_key=api_key)
                            prompt_text = get_autonomous_prompt(doc_type)
                            prompt = PromptTemplate(template=prompt_text, input_variables=["context", "text", "doc_type"])
                            final_prompt = prompt.format(context=contexto, text=raw_text[:60000], doc_type=doc_type)
                            
                            response = llm.invoke(final_prompt)
                            
                            # 3. Mostra e Gera Download
                            st.success("Análise Concluída!")
                            st.markdown(response.content)
                            
                            word_data = create_word_docx(response.content)
                            st.download_button("📥 Baixar Relatório (.docx)", word_data, "Auditoria.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                            
                            # Loga a operação
                            log_action(user["username"], "AUDITORIA", f"Doc: {doc_type}")

                    except Exception as e:
                        st.error(f"Erro técnico: {e}")

    # Outros módulos (Placeholders)
    else:
        st.button("⬅️ Voltar", on_click=lambda: st.session_state.update({"modulo_ativo": None}))
        st.title("🚧 Módulo em Configuração")
        st.info("Este módulo estará disponível assim que a configuração do seu pacote for finalizada.")

# --- SUPORTE ---
elif menu == "Suporte / Ajuda":
    st.title("Central de Suporte")
    
    tab1, tab2 = st.tabs(["📞 Contato", "🎓 Tutoriais"])
    
    with tab1:
        st.subheader("Fale com o Suporte Técnico")
        
        # [EDITAR] Mude o número abaixo para o seu WhatsApp
        whatsapp_link = "https://wa.me/5527999999999?text=Ol%C3%A1%2C%20preciso%20de%20ajuda%20no%20LiciGovtech"
        
        st.markdown(f"""
            <a href="{whatsapp_link}" target="_blank">
                <button style="background-color:#25D366; color:white; border:none; padding:15px 30px; border-radius:8px; font-size:16px; cursor:pointer;">
                    📱 Chamar no WhatsApp Agora
                </button>
            </a>
        """, unsafe_allow_html=True)
        
        st.divider()
        st.write("Ou abra um chamado via sistema:")
        msg = st.text_area("Descreva o problema:")
        if st.button("Enviar Chamado"):
            send_support_ticket(user["username"], msg)

    with tab2:
        st.subheader("Tutoriais de Uso")
        # [EDITAR] Aqui você coloca seus vídeos
        with st.expander("🎥 Como fazer uma Auditoria?"):
            st.write("Assista abaixo como auditar um Edital em 1 minuto.")
            # st.video("https://www.youtube.com/watch?v=SEU_VIDEO_AQUI") # Descomente e ponha o link
            st.info("(Vídeo tutorial em breve)")

# --- ADMIN ---
elif menu == "Admin":
    st.title("Painel Administrativo")
    st.subheader("Criar Novo Usuário")
    with st.form("new_user"):
        u_login = st.text_input("Login")
        u_pass = st.text_input("Senha")
        c1, c2, c3, c4, c5 = st.columns(5)
        p1 = c1.checkbox("Auditor", True)
        p2 = c2.checkbox("Gerador")
        p3 = c3.checkbox("Parecer")
        p4 = c4.checkbox("PCA")
        p5 = c5.checkbox("Recursos")
        if st.form_submit_button("Criar"):
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("INSERT INTO users (username, password, perm_auditor, perm_gerador, perm_parecer, perm_pca, perm_recursos) VALUES (%s, %s, %s, %s, %s, %s, %s)", 
                        (u_login, u_pass, p1, p2, p3, p4, p5))
            conn.commit()
            conn.close()
            st.success("Usuário Criado!")
    
    st.divider()
    st.subheader("Logs do Sistema")
    conn = get_db_connection()
    df_logs = conn.cursor()
    df_logs.execute("SELECT * FROM system_logs ORDER BY timestamp DESC LIMIT 10")
    st.table(df_logs.fetchall())
    conn.close()

